from ntpath import join
from secrets import choice
from docx import Document
import re
import easygui
from jinja2 import pass_eval_context
document = Document(r"D:\Users\Felipe\Desktop\Espaider\Robos\Editar word\Modelo - simulado 5.docx") 




NomeAluno = easygui.enterbox("Nome do Aluno/a: ")
print(NomeAluno)


###PEÇA GERAL###

#PEÇA
peca=easygui.ynbox("Acertou a peça?",choices=("[<F1>]Sim","[<F2>]Não"))
if peca==True:
    comentPeca='Parabéns! Você acertou a peça!'
else:
    comentPeca='Infelizmente você errou a peça, o que acarretaria na atribuição de nota zero e consequente reprovação. Contudo, com fins didáticos, iremos corrigir suas teses e argumentos apresentados.'

#ENDEREÇAMENTO
enderecamento=easygui.ynbox("Acertou endereçamento?",choices=("[<F1>]Sim","[<F2>]Não"))

if enderecamento==False:
    notaEnderecamento='0'
    comentEnderecamento='Infelizmente você errou o endereçamento.'
else:
    comentEnderecamento='Parabéns! Endereçamento feito corretamente!'
    notaEnderecamento=easygui.enterbox("Nota endereçamento: ")


#QUALIFICAÇÃO
partes=easygui.indexbox("Acertou qualificação?","Qualificação",("[<F1>]Sim","[<F2>]Parcialmente","[<F3>]Não"))

if partes==0:
    comentQualificacao='Parabéns! Qualificação das partes feita corretamente!'
elif partes==1:
    comentQualificacao='Pontuação atribuída parcialmente, pois você deixou de qualificar '
    erroQualificacao=easygui.enterbox("...deixou de qualificar:")
    comentQualificacao=comentQualificacao+erroQualificacao
else:
    comentQualificacao='Infelizmente você ou não qualificou, ou qualificou equivocadamente as partes.'
notaQualificacao=easygui.enterbox("Nota Qualificação")

#INTERPOSIÇÃO
partes=easygui.indexbox("Acertou interposição?","Interposição",("[<F1>]Sim","[<F2>]Parcialmente","[<F3>]Não"))

if partes==0:
    comentInterposicao='Parabéns! Pontuação atribuída integralmente!'
elif partes==1:
    comentInterposicao='Pontuação atribuída parcialmente, pois faltou '
    erroInterposicao=easygui.enterbox("...faltou:")
    comentInterposicao=comentInterposicao+erroInterposicao
else:
    comentInterposicao='Infelizmente você deixou de fazer a peça de interposição.'
notaInterposicao=easygui.enterbox("Nota Interposição")

#CONTRARRAZOES
partes=easygui.indexbox("Acertou contrarrazões?","Contrarrazões",("[<F1>]Sim","[<F2>]Parcialmente","[<F3>]Não"))

if partes==0:
    comentContrarrazoes='Parabéns! Pontuação atribuída integralmente!'
elif partes==1:
    comentContrarrazoes='Pontuação atribuída parcialmente, pois faltou '
    erroContrarrazoes=easygui.enterbox("...faltou:")
    comentContrarrazoes=comentContrarrazoes+erroContrarrazoes
else:
    comentContrarrazoes='Infelizmente você deixou de fazer o pedido de intimação para contrarrazões.'
notaContrarrazoes=easygui.enterbox("Nota contrarrazões")

#PREPARO
partes=easygui.indexbox("Acertou preparo?","Preparo",("[<F1>]Sim","[<F2>]Parcialmente","[<F3>]Não"))

if partes==0:
    comentPreparo='Parabéns! Pontuação atribuída integralmente!'
elif partes==1:
    comentPreparo='Pontuação atribuída parcialmente, pois faltou '
    erroPreparo=easygui.enterbox("...faltou:")
    comentPreparo=comentPreparo+erroPreparo
else:
    comentPreparo='Infelizmente você deixou de pedir a juntada das guias de preparo recursal.'
notaPreparo=easygui.enterbox("Nota preparo")

#LEGITIMIDADE
legitimidade=easygui.indexbox("Acertou legitimidade?","Legitimidade",("[<F1>]Sim","[<F2>]Parcialmente","[<F3>]Não"))

if legitimidade==0:
    comentLegitimidade='Parabéns! Pontuação atribuída integralmente!'
elif legitimidade==1:
    comentLegitimidade='Pontuação atribuída parcialmente, pois faltou '
    erroLegitimidade=easygui.enterbox("...pois faltou:")
    comentLegitimidade=comentLegitimidade+erroLegitimidade
else:
    comentLegitimidade='Pontuação não atribuída. Infelizmente você deixou de apresentar o tópico.'
notaLegitimidade=easygui.enterbox("Nota Legitimidade")


#COMPETÊNCIA
competencia=easygui.indexbox("Acertou Competência?","Competência",("[<F1>]Sim","[<F2>]Parcialmente","[<F3>]Não"))

if competencia==0:
    comentCompetencia='Parabéns! Pontuação atribuída integralmente!'
elif competencia==1:
    comentCompetencia='Pontuação atribuída parcialmente, pois faltou '
    erroCompetencia=easygui.enterbox("...pois faltou:")
    comentCompetencia=comentCompetencia+erroCompetencia
else:
    comentCompetencia='Pontuação não atribuída. Infelizmente você deixou de apresentar o tópico.'
notaCompetencia=easygui.enterbox("Nota Competência")


#CABIMENTO
cabimento=easygui.indexbox("Acertou Cabimento?","Cabimento",("[<F1>]Sim","[<F2>]Parcialmente","[<F3>]Não"))

if cabimento==0:
    comentCabimento='Parabéns! Pontuação atribuída integralmente!'
elif cabimento==1:
    comentCabimento='Pontuação atribuída parcialmente, pois faltou '
    erroCabimento=easygui.enterbox("...pois faltou:")
    comentCabimento=comentCabimento+erroCabimento
else:
    comentCabimento='Pontuação não atribuída. Infelizmente você deixou de apresentar o tópico.'
notaCabimento=easygui.enterbox("Nota Cabimento")


#FUNDAMENTOS
#FUNDAMENTO 1 - MORA
fundamento1=easygui.indexbox("Acertou Fundamento 1?","Fundamento 1",("[<F1>]Sim","[<F2>]Parcialmente","[<F3>]Não"))

if fundamento1==0:
    comentFundamento1='Parabéns! Pontuação atribuída integralmente!'
elif fundamento1==1:
    comentFundamento1='Pontuação atribuída parcialmente, pois faltou '
    erroFundamento1=easygui.enterbox("...pois faltou:")
    comentFundamento1=comentFundamento1+erroFundamento1
else:
    comentFundamento1='Pontuação não atribuída. Infelizmente você deixou de apresentar o tópico.'
notaFundamento1=easygui.enterbox("Nota Fundamento 1")

'''
#FUNDAMENTO 2 - DIREITO
fundamento2=easygui.indexbox("Acertou Fundamento 2?","Fundamento 2",("[<F1>]Sim","[<F2>]Parcialmente","[<F3>]Não"))

if fundamento2==0:
    comentFundamento2='Parabéns! Pontuação atribuída integralmente!'
elif fundamento2==1:
    comentFundamento2='Pontuação atribuída parcialmente, pois faltou '
    erroFundamento2=easygui.enterbox("...pois faltou:")
    comentFundamento2=comentFundamento2+erroFundamento2
else:
    comentFundamento2='Pontuação não atribuída. Infelizmente você deixou de apresentar o tópico.'
notaFundamento2=easygui.enterbox("Nota Fundamento 2")
'''

#PEDIDOS
#PEDIDO 1 - RECEBIMENTO
pedido1=easygui.indexbox("Acertou Pedido 1?","Pedido 1",("[<F1>]Sim","[<F2>]Parcialmente","[<F3>]Não"))

if pedido1==0:
    comentPedido1='Parabéns! Pontuação atribuída integralmente!'
elif pedido1==1:
    comentPedido1='Pontuação atribuída parcialmente, pois faltou '
    erroPedido1=easygui.enterbox("...pois faltou:")
    comentPedido1=comentPedido1+erroPedido1
else:
    comentPedido1='Pontuação não atribuída. Infelizmente você deixou de fazer o pedido.'
notaPedido1=easygui.enterbox("Nota Pedido 1")


#PEDIDO 2 - NOTIFICAÇÃO IMPETRADOS
pedido2=easygui.indexbox("Acertou Pedido 2?","Pedido 2",("[<F1>]Sim","[<F2>]Parcialmente","[<F3>]Não"))

if pedido2==0:
    comentPedido2='Parabéns! Pontuação atribuída integralmente!'
elif pedido2==1:
    comentPedido2='Pontuação atribuída parcialmente, pois faltou '
    erroPedido2=easygui.enterbox("...pois faltou:")
    comentPedido2=comentPedido2+erroPedido2
else:
    comentPedido2='Pontuação não atribuída. Infelizmente você deixou de fazer o pedido.'
notaPedido2=easygui.enterbox("Nota Pedido 2")


#PEDIDO 3 - NOTIFICAÇÃO REPRESENTAÇÃO
pedido3=easygui.indexbox("Acertou Pedido 3?","Pedido 3",("[<F1>]Sim","[<F2>]Parcialmente","[<F3>]Não"))

if pedido3==0:
    comentPedido3='Parabéns! Pontuação atribuída integralmente!'
elif pedido3==1:
    comentPedido3='Pontuação atribuída parcialmente, pois faltou '
    erroPedido3=easygui.enterbox("...pois faltou:")
    comentPedido3=comentPedido3+erroPedido3
else:
    comentPedido3='Pontuação não atribuída. Infelizmente você deixou de fazer o pedido.'
notaPedido3=easygui.enterbox("Nota Pedido 3")


#PEDIDO 4 - MP
pedido4=easygui.indexbox("Acertou Pedido 4?","Pedido 4",("[<F1>]Sim","[<F2>]Parcialmente","[<F3>]Não"))

if pedido4==0:
    comentPedido4='Parabéns! Pontuação atribuída integralmente!'
elif pedido4==1:
    comentPedido4='Pontuação atribuída parcialmente, pois faltou '
    erroPedido4=easygui.enterbox("...pois faltou:")
    comentPedido4=comentPedido4+erroPedido4
else:
    comentPedido4='Pontuação não atribuída. Infelizmente você deixou de fazer o pedido.'
notaPedido4=easygui.enterbox("Nota Pedido 4")

'''
#PEDIDO 5 - PROCEDÊNCIA
pedido5=easygui.indexbox("Acertou Pedido 5?","Pedido 5",("[<F1>]Sim","[<F2>]Parcialmente","[<F3>]Não"))

if pedido5==0:
    comentPedido5='Parabéns! Pontuação atribuída integralmente!'
elif pedido5==1:
    comentPedido5='Pontuação atribuída parcialmente, pois faltou '
    erroPedido5=easygui.enterbox("...pois faltou:")
    comentPedido5=comentPedido5+erroPedido5
else:
    comentPedido5='Pontuação não atribuída. Infelizmente você deixou de fazer o pedido.'
notaPedido5=easygui.enterbox("Nota Pedido 5")


#PEDIDO 6 - GARANTIA DE EFETIVIDADE
pedido6=easygui.indexbox("Acertou Pedido 6?","Pedido 6",("[<F1>]Sim","[<F2>]Parcialmente","[<F3>]Não"))

if pedido6==0:
    comentPedido6='Parabéns! Pontuação atribuída integralmente!'
elif pedido6==1:
    comentPedido6='Pontuação atribuída parcialmente, pois faltou '
    erroPedido6=easygui.enterbox("...pois faltou:")
    comentPedido6=comentPedido6+erroPedido6
else:
    comentPedido6='Pontuação não atribuída. Infelizmente você deixou de fazer o pedido.'
notaPedido6=easygui.enterbox("Nota Pedido 6")


#VALOR DA CAUSA
valor=easygui.indexbox("Acertou valor da causa?","Valor da Causa",("[<F1>]Sim","[<F2>]Parcialmente","[<F3>]Não"))

if valor==0:
    comentValor='Parabéns! Pontuação atribuída integralmente!'
elif valor==1:
    comentValor='Pontuação atribuída parcialmente, pois faltou '
    erroValor=easygui.enterbox("...pois faltou:")
    comentValor=comentValor+erroValor
else:
    comentValor='Pontuação não atribuída. Infelizmente você não apresentou o valor da causa.'
notaValor=easygui.enterbox("Nota valor da causa")
'''
#FECHAMENTO
fechamento=easygui.indexbox("Acertou Fechamento?","Fechamento",("[<F1>]Sim","[<F2>]Parcialmente","[<F3>]Não"))

if fechamento==0:
    comentFechamento='Parabéns! Pontuação atribuída integralmente!'
elif fechamento==1:
    comentFechamento='Pontuação atribuída parcialmente, pois faltou '
    erroFechamento=easygui.enterbox("...pois faltou:")
    comentFechamento=comentFechamento+erroFechamento
else:
    comentFechamento='Pontuação não atribuída. Infelizmente você deixou de fazer o fechamento.'
notaFechamento=easygui.enterbox("Nota Fechamento")

try:
    totalPeca=float(notaEnderecamento)+float(notaQualificacao)+float(notaLegitimidade)+float(notaInterposicao)+float(notaContrarrazoes)+float(notaPreparo)+float(notaCompetencia)+float(notaCabimento)+float(notaFundamento1)+float(notaPedido1)+float(notaPedido2)+float(notaPedido3)+float(notaPedido4)+float(notaFechamento)
    totalPeca=round(totalPeca,2)
    totalPeca=str(totalPeca)
    print(totalPeca)
except:
    print('Erro na soma das notas')
    totalPeca='XX'




###QUESTÕES###
#QUESTÃO 1A
questao1a=easygui.indexbox("Acertou Questão 1A?","Questão 1A",("[<F1>]Sim","[<F2>]Parcialmente","[<F3>]Não"))

if questao1a==0:
    comentQuestao1a='Parabéns! Desenvolvimento correto na resposta! Pontuação atribuída integralmente!'
elif questao1a==1:
    comentQuestao1a='Pontuação atribuída parcialmente, pois faltou '
    erroQuestao1a=easygui.enterbox("...pois faltou:")
    comentQuestao1a=comentQuestao1a+erroQuestao1a+'. Lembre-se sempre de deixar sua resposta completa'
else:
    comentQuestao1a='Pontuação não atribuída. Infelizmente você errou a resposta.'
notaQuestao1a=easygui.enterbox("Nota questão 1A")


#QUESTÃO 1B
questao1b=easygui.indexbox("Acertou Questão 1B?","Questão 1B",("[<F1>]Sim","[<F2>]Parcialmente","[<F3>]Não"))

if questao1b==0:
    comentQuestao1b='Parabéns! Desenvolvimento correto na resposta! Pontuação atribuída integralmente!'
elif questao1b==1:
    comentQuestao1b='Pontuação atribuída parcialmente, pois faltou '
    erroQuestao1b=easygui.enterbox("...pois faltou:")
    comentQuestao1b=comentQuestao1b+erroQuestao1b+'. Lembre-se sempre de deixar sua resposta completa'
else:
    comentQuestao1b='Pontuação não atribuída. Infelizmente você errou a resposta.'
notaQuestao1b=easygui.enterbox("Nota questão 1B")


#QUESTÃO 2A
questao2a=easygui.indexbox("Acertou Questão 2A?","Questão 2A",("[<F1>]Sim","[<F2>]Parcialmente","[<F3>]Não"))

if questao2a==0: 
    comentQuestao2a='Parabéns! Desenvolvimento correto na resposta! Pontuação atribuída integralmente!'
elif questao2a==1:
    comentQuestao2a='Pontuação atribuída parcialmente, pois faltou '
    erroQuestao2a=easygui.enterbox("...pois faltou:")
    comentQuestao2a=comentQuestao2a+erroQuestao2a+'. Lembre-se sempre de deixar sua resposta completa'
else:
    comentQuestao2a='Pontuação não atribuída. Infelizmente você errou a resposta.'
notaQuestao2a=easygui.enterbox("Nota questão 2A")


#QUESTÃO 2B
questao2b=easygui.indexbox("Acertou Questão 2B?","Questão 2B",("[<F1>]Sim","[<F2>]Parcialmente","[<F3>]Não"))

if questao2b==0:
    comentQuestao2b='Parabéns! Desenvolvimento correto na resposta! Pontuação atribuída integralmente!'
elif questao2b==1:
    comentQuestao2b='Pontuação atribuída parcialmente, pois faltou '
    erroQuestao2b=easygui.enterbox("...pois faltou:")
    comentQuestao2b=comentQuestao2b+erroQuestao2b+'. Lembre-se sempre de deixar sua resposta completa'
else:
    comentQuestao2b='Pontuação não atribuída. Infelizmente você errou a resposta.'
notaQuestao2b=easygui.enterbox("Nota questão 2B")


#QUESTÃO 3A
questao3a=easygui.indexbox("Acertou Questão 3A?","Questão 3A",("[<F1>]Sim","[<F2>]Parcialmente","[<F3>]Não"))

if questao3a==0:
    comentQuestao3a='Parabéns! Desenvolvimento correto na resposta! Pontuação atribuída integralmente!'
elif questao3a==1:
    comentQuestao3a='Pontuação atribuída parcialmente, pois faltou '
    erroQuestao3a=easygui.enterbox("...pois faltou:")
    comentQuestao3a=comentQuestao3a+erroQuestao3a+'. Lembre-se sempre de deixar sua resposta completa'
else:
    comentQuestao3a='Pontuação não atribuída. Infelizmente você errou a resposta.'
notaQuestao3a=easygui.enterbox("Nota questão 3A")


#QUESTÃO 3B
questao3b=easygui.indexbox("Acertou Questão 3B?","Questão 3B",("[<F1>]Sim","[<F2>]Parcialmente","[<F3>]Não"))

if questao3b==0:
    comentQuestao3b='Parabéns! Desenvolvimento correto na resposta! Pontuação atribuída integralmente!'
elif questao3b==1:
    comentQuestao3b='Pontuação atribuída parcialmente, pois faltou '
    erroQuestao3b=easygui.enterbox("...pois faltou:")
    comentQuestao3b=comentQuestao3b+erroQuestao3b+'. Lembre-se sempre de deixar sua resposta completa'
else:
    comentQuestao3b='Pontuação não atribuída. Infelizmente você errou a resposta.'
notaQuestao3b=easygui.enterbox("Nota questão 3B")


#QUESTÃO 4A
questao4a=easygui.indexbox("Acertou Questão 4A?","Questão 4A",("[<F1>]Sim","[<F2>]Parcialmente","[<F3>]Não"))

if questao4a==0:
    comentQuestao4a='Parabéns! Desenvolvimento correto na resposta! Pontuação atribuída integralmente!'
elif questao4a==1:
    comentQuestao4a='Pontuação atribuída parcialmente, pois faltou '
    erroQuestao4a=easygui.enterbox("...pois faltou:")
    comentQuestao4a=comentQuestao4a+erroQuestao4a+'. Lembre-se sempre de deixar sua resposta completa'
else:
    comentQuestao4a='Pontuação não atribuída. Infelizmente você errou a resposta.'
notaQuestao4a=easygui.enterbox("Nota questão 4A")


#QUESTÃO 4B
questao4b=easygui.indexbox("Acertou Questão 4B?","Questão 4B",("[<F1>]Sim","[<F2>]Parcialmente","[<F3>]Não"))

if questao4b==0:
    comentQuestao4b='Parabéns! Desenvolvimento correto na resposta! Pontuação atribuída integralmente!'
elif questao4b==1:
    comentQuestao4b='Pontuação atribuída parcialmente, pois faltou '
    erroQuestao4b=easygui.enterbox("...pois faltou:")
    comentQuestao4b=comentQuestao4b+erroQuestao4b+'. Lembre-se sempre de deixar sua resposta completa'
else:
    comentQuestao4b='Pontuação não atribuída. Infelizmente você errou a resposta.'
notaQuestao4b=easygui.enterbox("Nota questão 4B")

'''
#OBSERVAÇÕES
observacoes=easygui.multchoicebox("Alguma observação?","Observações",("Rasuras","Margens","Não inventar dados","Letra","Nenhuma"))

rasuras='Cuidado com as rasuras! No caso de escrita equivocada, apenas faça um traço simples e continue escrevendo ao lado.'
margens='Cuidado com a escrita fora das margens e linhas, sob pena de não obter correção no referido ponto.'
inventaDados='Cuidado: Não invente dados que o enunciado não aborda. A FGV poderá considerar identificação de peça e zerá-la. '
letra='Cuidado: Sua letra não está totalmente legível, tente melhorá-la para que o examinador não desconte nenhum ponto porque não entendeu o que estava escrito.'
'''


#PARÁGRAFOS
for paragraph in document.paragraphs:
    if '$NomeAluno' in paragraph.text:
        texto = paragraph.text
        texto = texto.replace('$NomeAluno',NomeAluno)
        paragraph.text = texto
    '''
    if '$rasura' in paragraph.text:
        texto = paragraph.text
        if "Rasuras" in observacoes:
            texto = texto.replace('$rasuras',rasuras)
        else:
            texto = texto.replace('$rasura','')
        paragraph.text = texto
    if '$margens' in paragraph.text:
        texto = paragraph.text
        if "Margens" in observacoes:
            texto = texto.replace('$margens',margens)
        else:
            texto = texto.replace('$margens','')
        paragraph.text = texto
    if '$inventaDados' in paragraph.text:
        texto = paragraph.text
        if "Não inventar dados" in observacoes:
            texto = texto.replace('$inventaDados',inventaDados)
        else:
            texto = texto.replace('$inventaDados','')
        paragraph.text = texto
    if '$letra' in paragraph.text:
        texto = paragraph.text
        if "Letra" in observacoes:
            texto = texto.replace('$letra',letra)
        else:
            texto = texto.replace('$letra','')
        paragraph.text = texto
    '''
    if '$comentPeca' in paragraph.text:
        texto = paragraph.text
        texto = texto.replace('$comentPeca',comentPeca)
        paragraph.text = texto
    if '$comentEnderecamento' in paragraph.text:
        texto = paragraph.text
        texto = texto.replace('$comentEnderecamento',comentEnderecamento)
        paragraph.text = texto
    if '$comentQualificacao' in paragraph.text:
        texto = paragraph.text
        texto = texto.replace('$comentQualificacao',comentQualificacao)
        paragraph.text = texto
    if '$comentInterposicao' in paragraph.text:
        texto = paragraph.text
        texto = texto.replace('$comentInterposicao',comentInterposicao)
        paragraph.text = texto
    if '$comentContrarrazoes' in paragraph.text:
        texto = paragraph.text
        texto = texto.replace('$comentContrarrazoes',comentContrarrazoes)
        paragraph.text = texto
    if '$comentPreparo' in paragraph.text:
        texto = paragraph.text
        texto = texto.replace('$comentPreparo',comentPreparo)
        paragraph.text = texto
    if '$comentLegitimidade' in paragraph.text:
        texto = paragraph.text
        texto = texto.replace('$comentLegitimidade',comentLegitimidade)
        paragraph.text = texto
    if '$comentCompetencia' in paragraph.text:
        texto = paragraph.text
        texto = texto.replace('$comentCompetencia',comentCompetencia)
        paragraph.text = texto
    if '$comentCabimento' in paragraph.text:
        texto = paragraph.text
        texto = texto.replace('$comentCabimento',comentCabimento)
        paragraph.text = texto
    if '$comentFundamento1' in paragraph.text:
        texto = paragraph.text
        texto = texto.replace('$comentFundamento1',comentFundamento1)
        paragraph.text = texto
    '''
    if '$comentFundamento2' in paragraph.text:
        texto = paragraph.text
        texto = texto.replace('$comentFundamento2',comentFundamento2)
        paragraph.text = texto
    '''
    if '$comentPedido1' in paragraph.text:
        texto = paragraph.text
        texto = texto.replace('$comentPedido1',comentPedido1)
        paragraph.text = texto
    if '$comentPedido2' in paragraph.text:
        texto = paragraph.text
        texto = texto.replace('$comentPedido2',comentPedido2)
        paragraph.text = texto
    if '$comentPedido3' in paragraph.text:
        texto = paragraph.text
        texto = texto.replace('$comentPedido3',comentPedido3)
        paragraph.text = texto
    if '$comentPedido4' in paragraph.text:
        texto = paragraph.text
        texto = texto.replace('$comentPedido4',comentPedido4)
        paragraph.text = texto
    '''
    if '$comentPedido5' in paragraph.text:
        texto = paragraph.text
        texto = texto.replace('$comentPedido5',comentPedido5)
        paragraph.text = texto
    if '$comentPedido6' in paragraph.text:
        texto = paragraph.text
        texto = texto.replace('$comentPedido6',comentPedido6)
        paragraph.text = texto
    if '$comentValor' in paragraph.text:
        texto = paragraph.text
        texto = texto.replace('$comentValor',comentValor)
        paragraph.text = texto
    '''
    if '$comentFechamento' in paragraph.text:
        texto = paragraph.text
        texto = texto.replace('$comentFechamento',comentFechamento)
        paragraph.text = texto
        
    


#TABELAS
for table in document.tables:
    for row in table.rows:
        for cell in row.cells:    
            if '$notaEnderecamento' in cell.text:
                texto = cell.text
                texto = texto.replace('$notaEnderecamento',notaEnderecamento)
                cell.text = texto
            if '$notaQualificacao' in cell.text:
                texto = cell.text
                texto = texto.replace('$notaQualificacao',notaQualificacao)
                cell.text = texto
            if '$notaInterposicao' in cell.text:
                texto = cell.text
                texto = texto.replace('$notaInterposicao',notaInterposicao)
                cell.text = texto
            if '$notaContrarrazoes' in cell.text:
                texto = cell.text
                texto = texto.replace('$notaContrarrazoes',notaContrarrazoes)
                cell.text = texto
            if '$notaPreparo' in cell.text:
                texto = cell.text
                texto = texto.replace('$notaPreparo',notaPreparo)
                cell.text = texto
            if '$notaLegitimidade' in cell.text:
                texto = cell.text
                texto = texto.replace('$notaLegitimidade',notaLegitimidade)
                cell.text = texto
            if '$notaCompetencia' in cell.text:
                texto = cell.text
                texto = texto.replace('$notaCompetencia',notaCompetencia)
                cell.text = texto
            if '$notaCabimento' in cell.text:
                texto = cell.text
                texto = texto.replace('$notaCabimento',notaCabimento)
                cell.text = texto
            if '$notaFundamento1' in cell.text:
                texto = cell.text
                texto = texto.replace('$notaFundamento1',notaFundamento1)
                cell.text = texto
            '''    
            if '$notaFundamento2' in cell.text:
                texto = cell.text
                texto = texto.replace('$notaFundamento2',notaFundamento2)
                cell.text = texto
            '''
            if '$notaPedido1' in cell.text:
                texto = cell.text
                texto = texto.replace('$notaPedido1',notaPedido1)
                cell.text = texto
            if '$notaPedido2' in cell.text:
                texto = cell.text
                texto = texto.replace('$notaPedido2',notaPedido2)
                cell.text = texto
            if '$notaPedido3' in cell.text:
                texto = cell.text
                texto = texto.replace('$notaPedido3',notaPedido3)
                cell.text = texto
            if '$notaPedido4' in cell.text:
                texto = cell.text
                texto = texto.replace('$notaPedido4',notaPedido4)
                cell.text = texto
            '''
            if '$notaPedido5' in cell.text:
                texto = cell.text
                texto = texto.replace('$notaPedido5',notaPedido5)
                cell.text = texto
            if '$notaPedido6' in cell.text:
                texto = cell.text
                texto = texto.replace('$notaPedido6',notaPedido6)
                cell.text = texto
            if '$notaValor' in cell.text:
                texto = cell.text
                texto = texto.replace('$notaValor',notaValor)
                cell.text = texto
            '''
            if '$notaFechamento' in cell.text:
                texto = cell.text
                texto = texto.replace('$notaFechamento',notaFechamento)
                cell.text = texto
            if '$totalPeca' in cell.text:
                texto = cell.text
                texto = texto.replace('$totalPeca',totalPeca)
                cell.text = texto
            if '$notaQuestao1a' in cell.text:
                texto = cell.text
                texto = texto.replace('$notaQuestao1a',notaQuestao1a)
                cell.text = texto
            if '$notaQuestao1b' in cell.text:
                texto = cell.text
                texto = texto.replace('$notaQuestao1b',notaQuestao1b)
                cell.text = texto
            if '$comentQuestao1a' in cell.text:
                texto = cell.text
                texto = texto.replace('$comentQuestao1a',comentQuestao1a)
                cell.text = texto
            if '$comentQuestao1b' in cell.text:
                texto = cell.text
                texto = texto.replace('$comentQuestao1b',comentQuestao1b)
                cell.text = texto
            if '$notaQuestao2a' in cell.text:
                texto = cell.text
                texto = texto.replace('$notaQuestao2a',notaQuestao2a)
                cell.text = texto
            if '$notaQuestao2b' in cell.text:
                texto = cell.text
                texto = texto.replace('$notaQuestao2b',notaQuestao2b)
                cell.text = texto
            if '$comentQuestao2a' in cell.text:
                texto = cell.text
                texto = texto.replace('$comentQuestao2a',comentQuestao2a)
                cell.text = texto
            if '$comentQuestao2b' in cell.text:
                texto = cell.text
                texto = texto.replace('$comentQuestao2b',comentQuestao2b)
                cell.text = texto
            if '$notaQuestao3a' in cell.text:
                texto = cell.text
                texto = texto.replace('$notaQuestao3a',notaQuestao3a)
                cell.text = texto
            if '$notaQuestao3b' in cell.text:
                texto = cell.text
                texto = texto.replace('$notaQuestao3b',notaQuestao3b)
                cell.text = texto
            if '$comentQuestao3a' in cell.text:
                texto = cell.text
                texto = texto.replace('$comentQuestao3a',comentQuestao3a)
                cell.text = texto
            if '$comentQuestao3b' in cell.text:
                texto = cell.text
                texto = texto.replace('$comentQuestao3b',comentQuestao3b)
                cell.text = texto
            if '$notaQuestao4a' in cell.text:
                texto = cell.text
                texto = texto.replace('$notaQuestao4a',notaQuestao4a)
                cell.text = texto
            if '$notaQuestao4b' in cell.text:
                texto = cell.text
                texto = texto.replace('$notaQuestao4b',notaQuestao4b)
                cell.text = texto
            if '$comentQuestao4a' in cell.text:
                texto = cell.text
                texto = texto.replace('$comentQuestao4a',comentQuestao4a)
                cell.text = texto
            if '$comentQuestao4b' in cell.text:
                texto = cell.text
                texto = texto.replace('$comentQuestao4b',comentQuestao4b)
                cell.text = texto


document.save(f'padrao-de-resposta-{NomeAluno}.docx')